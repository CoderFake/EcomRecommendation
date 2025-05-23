import requests
from django.contrib.sites.shortcuts import get_current_site
from django.db.models import Count, Sum
from datetime import datetime, timedelta
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
import requests
import json
import base64
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
from cryptography.hazmat.primitives import padding
from cryptography.hazmat.backends import default_backend
from decouple import config
from django.db.models.functions import ExtractHour
from django.forms import model_to_dict
from django.utils.dateparse import parse_date, parse_datetime
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from requests import Response
import openpyxl
from EcomRecomendation import settings
from accounts.models import Account, UserProfile, EventUser
from accounts.views import upload_image_to_cloudflare
from django.shortcuts import redirect
from category.models import CategoryMain, SubCategory
from category.forms import SubCategoryForm, CategoryMainForm
from store.models import Product, Variation, VariationManager, FolderEvent, Wishlist
from carts.forms import ProductForm
from store.forms import variationForm
from orders.forms import OrderForm, OrderUpdateForm
from django.http import HttpResponse, JsonResponse
from orders.models import Order, OrderProduct, Payment
from django.contrib import messages
import json
import pandas as pd
from django.utils import timezone
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from datetime import timedelta
from django.http import HttpResponse
from django.utils.dateparse import parse_datetime
from django.contrib.auth.decorators import login_required
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
import os
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger

def get_locations(id, parent_id):
    api_url = "https://member.lazada.vn/locationtree/api/getSubAddressList?countryCode=VN"
    if parent_id:
        api_url += f"&addressId={parent_id}"
    response = requests.get(api_url)
    locations = response.json().get('module', [])
    for location in locations:
        if location.get('id') == id:
            return location.get('name')
    return None


def update_dash(request):
    if request.method == "POST":
        timezone.activate(settings.TIME_ZONE)
        today = timezone.now().date()
        yesterday = today - timedelta(days=1)


        users_logged_in_today = Account.objects.filter(last_login__date__gte=today).count()
        users_logged_in_not_today = Account.objects.filter(last_login__date__lt=today, is_login=True).count()
        daily_signins = users_logged_in_today + users_logged_in_not_today
        daily_signup = Account.objects.filter(date_joined__date__gte=today).count()
        daily_order = Order.objects.filter(created_at__date__gte=today).count()
        orders = Order.objects.filter(created_at__gte=today, is_ordered=True)
        daily_total_order = sum(order.payment.amount_paid for order in orders)


        users_logged_in_yesterday = Account.objects.filter(last_login__date__gte=yesterday,
                                                           last_login__date__lt=today).count()
        users_logged_in_not_yesterday = Account.objects.filter(last_login__date__lt=yesterday, is_login=True).count()
        yesterday_signins = users_logged_in_yesterday + users_logged_in_not_yesterday
        yesterday_signup = Account.objects.filter(date_joined__date__gte=yesterday, date_joined__date__lt=today).count()
        yesterday_order = Order.objects.filter(created_at__date__gte=yesterday, created_at__date__lt=today).count()
        yesterday_orders = Order.objects.filter(created_at__gte=yesterday, created_at__lt=today, is_ordered=True)
        yesterday_total_order = sum(order.payment.amount_paid for order in yesterday_orders)

        def calculate_percentage_change(today_value, yesterday_value):
            if yesterday_value == 0:
                return 100 if today_value > 0 else 0
            return round(((today_value - yesterday_value) / yesterday_value) * 100, 2)

        signin_change = calculate_percentage_change(daily_signins, yesterday_signins)
        signup_change = calculate_percentage_change(daily_signup, yesterday_signup)
        order_change = calculate_percentage_change(daily_order, yesterday_order)
        total_order_change = calculate_percentage_change(daily_total_order, yesterday_total_order)

        context = {
            "user_signins": daily_signins,
            "signin_change": signin_change,
            "daily_signups": daily_signup,
            "signup_change": signup_change,
            "daily_order": daily_order,
            "order_change": order_change,
            "daily_total_order": round(daily_total_order, 2),
            "total_order_change": total_order_change
        }

        return JsonResponse(context, safe=False)


@login_required(login_url='login')
def index(request):
    if request.user.is_superuser:
        return render(request, 'adminApp/dashboard.html')
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def user_infor(request):
    if request.method == 'POST':
        if request.user.is_superuser:
            try:
                user_profile = UserProfile.objects.get(user=request.user)
                data = {
                    'date_of_birth': user_profile.date_of_birth.strftime(
                        "%Y-%m-%d") if user_profile.date_of_birth else None,
                    'sex': user_profile.sex,
                    'road': user_profile.road,
                    'ward': user_profile.ward,
                    'district': user_profile.district,
                    'city': user_profile.city,
                    'profile_picture': user_profile.profile_picture,
                    'bio': user_profile.bio
                }
                return JsonResponse(data)
            except UserProfile.DoesNotExist:
                return JsonResponse({'error': 'User profile not found'}, status=404)
        else:
            return JsonResponse({'error': 'You are not authorized to view this page'}, status=403)
    else:
        return JsonResponse({'error': 'Invalid request, only POST method is allowed'}, status=400)


@login_required(login_url='login')
def user_list(request):
    if request.user.is_superuser:
        profile_list = UserProfile.objects.all().order_by('-user__date_joined')
        paginator = Paginator(profile_list, 10) 

        page_number = request.GET.get('page')
        try:
            page_obj = paginator.get_page(page_number)
        except PageNotAnInteger:
            page_obj = paginator.page(1)
        except EmptyPage:
            page_obj = paginator.page(paginator.num_pages)

        profile_user = UserProfile.objects.get(user__pk=request.user.id)
        context = {
            'page_obj': page_obj, 
            'profile_user': profile_user
        }
        return render(request, 'adminApp/users/user_list.html', context)
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def user_profile(request, pk):
    if request.user.is_superuser:
        if request.method == 'POST':
            user = request.user
            profile = UserProfile.objects.get(user__pk=pk)
            user.full_name = request.POST.get('full_name')
            user.username = request.POST.get('user_name')
            user.phone_number = request.POST.get('phone_number')
            user.save()

            date_of_birth = request.POST.get('date_of_birth')
            if date_of_birth:
                try:
                    datetime.strptime(date_of_birth, '%Y-%m-%d')
                    profile.date_of_birth = date_of_birth
                except ValueError:
                    print("Incorrect date format")

            profile.sex = request.POST.get('sex')
            profile.road = request.POST.get('road')
            profile.ward = request.POST.get('ward')
            profile.district = request.POST.get('district')
            profile.city = request.POST.get('city')
            profile.bio = request.POST.get('bio')

            # Xử lý file ảnh
            if 'profile_picture' in request.FILES:
                try:
                    image_url = upload_image_to_cloudflare(request.FILES['profile_picture'], request)
                    if image_url is not None:
                        profile.profile_picture = image_url
                except Exception as e:
                    messages.error(request, f"Failed to upload image: {str(e)}")

            profile.save()

            old_password = request.POST.get('old_password')
            new_password = request.POST.get('new_password')
            confirm_password = request.POST.get('confirm_password')

            if old_password and new_password and confirm_password:
                if new_password == confirm_password:
                    if user.check_password(old_password):
                        user.set_password(new_password)
                        user.save()
                    else:
                        messages.error(request, "Old password is incorrect.")
                else:
                    messages.error(request, "New password and confirm password do not match.")

            messages.success(request, 'Profile updated successfully.')
            return redirect('user_profile', pk=profile.pk)

        profile = UserProfile.objects.get(user__pk=pk)
        context = {
            'profile': profile,
        }
        return render(request, 'adminApp/users/user_profile.html', context)
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def user_delete(request, pk):
    if request.user.is_superuser:
        user = Account.objects.get(pk=pk)
        user.delete()
        return redirect('user_list')
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def user_block(request, pk):
    if request.user.is_superuser:
        user = Account.objects.get(pk=pk)
        user.is_active = False
        user.save()
        return redirect('user_list')
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def user_unblock(request, pk):
    if request.user.is_superuser:
        user = Account.objects.get(pk=pk)
        user.is_active = True
        user.save()
        return redirect('user_list')
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def get_login_frequency(request):
    if request.method == 'POST' and request.user.is_superuser:
        try:
            data = json.loads(request.body)
            start_date = parse_datetime(data['start_date'])
            end_date = parse_datetime(data['end_date'])
            end_date = end_date + timedelta(days=1) - timedelta(seconds=1)
        except (ValueError, KeyError):
            return JsonResponse({'error': 'Invalid data'}, status=400)

        events = EventUser.objects.filter(
            event_timestamp__range=[start_date, end_date],
            event_type='login'
        ).annotate(hour=ExtractHour('event_timestamp')).values('hour').annotate(count=Count('id')).order_by('hour')

        data = [0] * 12
        for event in events:
            hour_index = event['hour'] // 2
            if hour_index < 12:
                data[hour_index] += event['count']

        return JsonResponse({'data': data})
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def get_user_acquisition(request):
    if request.method == 'POST' and request.user.is_superuser:
        data = json.loads(request.body)
        start_date = datetime.strptime(data['start_date'], '%Y-%m-%d')
        end_date = datetime.strptime(data['end_date'], '%Y-%m-%d') + timedelta(days=1)

        events = EventUser.objects.filter(
            event_timestamp__range=[start_date, end_date],
            event_type__in=['view', 'cart', 'pay']  
        ).annotate(hour=ExtractHour('event_timestamp')).values('hour', 'event_type').annotate(count=Count('id')).order_by('hour')

        result = {
            'View': [0]*12, 'Cart': [0]*12, 'Pay': [0]*12
        }

        for event in events:
            hour_index = event['hour'] // 2
            event_type = event['event_type'].capitalize()
            if event_type in result:
                result[event_type][hour_index] += 1

        return JsonResponse(result)
    else:
        return HttpResponse('You are not authorized to view this page', status=403)

@login_required(login_url='login')
def main_category(request):
    if request.user.is_superuser:
        if request.method == 'POST':
            form = CategoryMainForm(request.POST, request.FILES)
            if form.is_valid():
                form.save()
                messages.success(request, 'Main category added successfully.')
                return redirect('main_category')
            else:
                 messages.error(request, 'Failed to add main category. Please check the form.')

        main_list = CategoryMain.objects.all().order_by('category_name')
        paginator = Paginator(main_list, 10)

        page_number = request.GET.get('page')
        try:
            page_obj = paginator.get_page(page_number)
        except PageNotAnInteger:
            page_obj = paginator.page(1)
        except EmptyPage:
            page_obj = paginator.page(paginator.num_pages)

        data = SubCategory.objects.all() # Keep this for potential use in the form/template
        form = CategoryMainForm()
        context = {
            'page_obj': page_obj, # Pass the page object for main categories
            'sub': data,
            'form': form,
        }
        return render(request, 'adminApp/Category/MainCategory.html', context)
    return HttpResponse('You are not authorized to view this page')


def main_category_edit(request, pk):
    if request.user.is_superuser:
        main = CategoryMain.objects.get(pk=pk)
        data = SubCategory.objects.all()
        form = CategoryMainForm(instance=main)
        mains = CategoryMain.objects.all()
        if request.method == 'POST':
            form = CategoryMainForm(request.POST, instance=main)
            if form.is_valid():
                form.save()
                return redirect('main_category')
            else:
                print(form.errors)

        context = {
            'main': main,
            'edit_form': form,
            'sub': data,
            'mains': mains,
        }
        return render(request, 'adminApp/Category/MainCategoryedit.html', context)

    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def main_category_delete(request, pk):
    if request.user.is_superuser:
        main = CategoryMain.objects.get(pk=pk)
        main.delete()
        return redirect('main_category')
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def sub_category(request):
    if request.user.is_superuser:
        if request.method == 'POST':
            form = SubCategoryForm(request.POST, request.FILES)
            if form.is_valid():
                form.save()
                messages.success(request, 'Sub category added successfully.')
                return redirect('sub_category')
            else:
                messages.error(request, 'Failed to add sub category. Please check the form.')

        sub_list = SubCategory.objects.all().order_by('category__category_name', 'sub_category_name')
        paginator = Paginator(sub_list, 10)

        page_number = request.GET.get('page')
        try:
            page_obj = paginator.get_page(page_number)
        except PageNotAnInteger:
            page_obj = paginator.page(1)
        except EmptyPage:
            page_obj = paginator.page(paginator.num_pages)

        form = SubCategoryForm()
        context = {
            'page_obj': page_obj,
            'form': form,
        }
        return render(request, 'adminApp/Category/SubCategory.html', context)
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def sub_category_edit(request, pk):
    sub = SubCategory.objects.get(pk=pk)
    form = SubCategoryForm(instance=sub)
    if request.user.is_superuser:
        if request.method == 'POST':
            form = SubCategoryForm(request.POST, request.FILES, instance=sub)
            if form.is_valid():
                form.save()
                return redirect('sub_category')
            else:
                print(form.errors)

        context = {
            'sub': sub,
            'edit_form': form,
        }
        return render(request, 'adminApp/Category/SubCategoryedit.html', context)
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def sub_category_delete(request, pk):
    if request.user.is_superuser:
        sub = SubCategory.objects.get(pk=pk)
        sub.delete()
        return redirect('sub_category')
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def product_list(request):
    if request.user.is_superuser:
        product_list_all = Product.objects.all().order_by('-created_date') 
        paginator = Paginator(product_list_all, 10) 

        page_number = request.GET.get('page')
        try:
            page_obj = paginator.get_page(page_number)
        except PageNotAnInteger:
            page_obj = paginator.page(1)
        except EmptyPage:
            page_obj = paginator.page(paginator.num_pages)

        addProductForm = ProductForm() 
        if request.method == 'POST':
            form = ProductForm(request.POST, request.FILES)
            if form.is_valid():
                form.save()
                messages.success(request, 'Product added successfully.')
                return redirect('product_list')
            else:
                messages.error(request, 'Failed to add product. Please check the form.')
                addProductForm = form

        context = {
            'page_obj': page_obj,
            'addProduct': addProductForm,
        }
        return render(request, 'adminApp/Products/product_list.html', context)
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def product_edit(request, pk):
    if request.user.is_superuser:
        Allproduct = Product.objects.all()
        product = Product.objects.get(pk=pk)
        form = ProductForm(instance=product)
        if request.method == 'POST':
            form = ProductForm(request.POST, request.FILES, instance=product)
            if form.is_valid():
                form.save()
                return redirect('product_list')
            else:
                print(form.errors)

        context = {
            'product': Allproduct,
            'product_editing': product,
            'edit_form': form,
        }
        return render(request, 'adminApp/Products/product_edit.html', context)
    else:
        return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def product_delete(request, pk):
    if request.user.is_superuser:
        product = Product.objects.get(pk=pk)
        product.delete()
        return redirect('product_list')
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def add_variations(request):
    if request.user.is_superuser:
        variation_list = Variation.objects.all().order_by('Product__product_name', 'Variation_category', 'Variation_value')
        paginator = Paginator(variation_list, 10) 

        page_number = request.GET.get('page')
        try:
            page_obj = paginator.get_page(page_number)
        except PageNotAnInteger:
            page_obj = paginator.page(1)
        except EmptyPage:
            page_obj = paginator.page(paginator.num_pages)

        form = variationForm() 
        if request.method == 'POST':
            add_form = variationForm(request.POST, request.FILES)
            if add_form.is_valid():
                add_form.save()
                messages.success(request, 'Variation added successfully.')
                return redirect('add_variations')
            else:
                 messages.error(request, 'Failed to add variation. Please check the form.')
                 form = add_form 

        context = {
            'page_obj': page_obj, 
            'form': form,
        }
       
        return render(request, 'adminApp/Variations/add_variations.html', context)
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def edit_variations(request, pk):
    if request.user.is_superuser:
        existing_variations = Variation.objects.all()
        variation = Variation.objects.get(pk=pk)
        form = variationForm(instance=variation)
        if request.method == 'POST':
            form = variationForm(request.POST, request.FILES, instance=variation)
            if form.is_valid():
                form.save()
                return redirect('add_variations')
            else:
                print(form.errors)

        context = {
            'existing_variations': existing_variations,
            'variation_editing': variation,
            'form': form,
        }
        return render(request, 'adminApp/Variations/edit_variations.html', context)
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def delete_variations(request, pk):
    if request.user.is_superuser:
        variation = Variation.objects.get(pk=pk)
        variation.delete()
        return redirect('add_variations')
    return HttpResponse('You are not authorized to view this page')


MODEL_DIR = os.path.join(settings.BASE_DIR, 'static', 'adminapp', 'assets', 'img', 'logo')
DOC_DIR = os.path.join(settings.BASE_DIR, 'static', 'doc')

def set_cell_border(cell, **kwargs):

    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    for edge in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            element = OxmlElement(f"w:{edge}")
            element.set(qn('w:val'), edge_data.get("val", "single"))
            element.set(qn('w:sz'), str(edge_data.get("sz", 12)))
            element.set(qn('w:space'), str(edge_data.get("space", 0)))
            element.set(qn('w:color'), edge_data.get("color", "000000"))
            tcPr.append(element)


@login_required(login_url='login')
def download_order_report(request):
    if request.method == 'POST' and request.user.is_superuser:
        data = json.loads(request.body)
        start_date = parse_datetime(data['start_date'])
        end_date = parse_datetime(data['end_date'])
        end_date = end_date + timedelta(days=1) - timedelta(seconds=1)
        orders = Order.objects.filter(
            created_at__gte=start_date,
            created_at__lte=end_date
        ).values(
            'id', 'order_number', 'created_at', 'user__full_name', 'user__email',
            'order_total', 'payment__status', 'order_status', 'phone', 'road', 'ward', 'district', 'city'
        )

        # Create a Document
        doc = Document()

        # Set landscape orientation
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # Add title and logo in the same cell
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        paragraph = cell.paragraphs[0]

        # Add logo
        run = paragraph.add_run()
        run.add_picture(
            os.path.join(settings.BASE_DIR, 'static', 'adminapp', 'assets', 'img', 'logo', 'ec-site-logo.png'),
            width=Inches(1.0))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        store_name = cell.add_paragraph()
        run = store_name.add_run("EKKA Store")
        run.font.size = Pt(48)
        run.bold = True
        store_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        store_name.paragraph_format.space_before = Pt(12)

        title_para = doc.add_heading('Orders Overview Report', level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.size = Pt(22)

        date_para = doc.add_paragraph(f"From: {start_date.strftime('%Y-%m-%d')} To: {end_date.strftime('%Y-%m-%d')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.style = doc.styles['Normal']
        for run in date_para.runs:
            run.font.size = Pt(18)

        # Table
        table = doc.add_table(rows=1, cols=10)
        hdr_cells = table.rows[0].cells
        headers = ["ID", "Order Number", "Full Name", "Phone Number", "Email", "Address",
                   "Created At", "Payment Status", "Order Status", "Total ($)"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            set_cell_border(
                hdr_cells[i],
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            )

        # Insert order data
        for order in orders:
            row_cells = table.add_row().cells
            row_cells[0].text = str(order['id'])
            row_cells[1].text = order['order_number']
            row_cells[2].text = order['user__full_name']
            row_cells[3].text = order['phone']
            row_cells[4].text = order['user__email']
            row_cells[5].text = f"{order['road']}, {get_locations(order['ward'], order['district']) }, {get_locations(order['district'], order['city'])}, {get_locations(order['city'],'')}"
            row_cells[6].text = order['created_at'].strftime('%Y-%m-%d %H:%M:%S')
            row_cells[7].text = order['payment__status']
            row_cells[8].text = order['order_status']
            row_cells[9].text = str(order['order_total'])
            for cell in row_cells:
                set_cell_border(
                    cell,
                    top={"sz": 6, "val": "single", "color": "000000"},
                    bottom={"sz": 6, "val": "single", "color": "000000"},
                    left={"sz": 6, "val": "single", "color": "000000"},
                    right={"sz": 6, "val": "single", "color": "000000"},
                )

        # Save the document
        if not os.path.exists(DOC_DIR):
            os.makedirs(DOC_DIR)
        doc_file = os.path.join(DOC_DIR, 'orders_report.docx')
        doc.save(doc_file)

        # Read the file content and return as response
        with open(doc_file, 'rb') as f:
            response = HttpResponse(f.read(),
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response[
                'Content-Disposition'] = f'attachment; filename="orders_report_{start_date.strftime("%Y%m%d")}_to_{end_date.strftime("%Y%m%d")}.docx"'
            return response

    return HttpResponse("Invalid request", status=400)

@login_required(login_url='login')
def export_payments_to_excel(request):
    if request.method == 'POST' and request.user.is_superuser:
        try:
            data = json.loads(request.body)
            start_date = data.get('start_date')
            end_date = data.get('end_date')

            if not start_date or not end_date:
                return JsonResponse({'error': 'Missing start_date or end_date'}, status=400)

            start_date = datetime.strptime(start_date, '%Y-%m-%d')
            end_date = datetime.strptime(end_date, '%Y-%m-%d') + timedelta(days=1)

        except ValueError as e:
            return JsonResponse({'error': str(e)}, status=400)

        totals = {}

        payments = Payment.objects.filter(
            order__created_at__range=[start_date, end_date],
            order__is_ordered=True
        ).values('order__city', 'amount_paid')

        for payment in payments:
            city_name = payment['order__city']
            if city_name:
                if city_name not in totals:
                    totals[city_name] = 0
                totals[city_name] += payment['amount_paid']

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Payments by City"

        ws.append(["City", "Revenue"])

        for city, total in totals.items():
            ws.append([get_locations(city, ''), total])

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="payments_by_city.xlsx"'

        wb.save(response)

        return response

    return HttpResponse('You are not authorized to view this page', status=403)


def get_order_stats(request):
    if request.method == 'POST' and request.user.is_superuser:
        data = json.loads(request.body)
        start_date = parse_date(data.get('start_date'))
        end_date = parse_date(data.get('end_date'))
        period = data.get('period')

        orders = Order.objects.filter(created_at__date__range=[start_date, end_date])

        response_data = {
            'labels': [],
            'accepted': [],
            'ready_to_ship': [],
            'on_shipping': [],
            'delivered': [],
            'cancelled': [],
            'returned': []
        }

        if period == 'hour':
            for hour in range(24):
                start_time = datetime.combine(start_date, datetime.min.time()) + timedelta(hours=hour)
                end_time = start_time + timedelta(hours=1)
                response_data['labels'].append(f"{hour:02d}:00-{hour + 1:02d}:00")
                response_data['accepted'].append(
                    orders.filter(created_at__time__range=[start_time.time(), end_time.time()],
                                  order_status='Accepted').count())
                response_data['ready_to_ship'].append(
                    orders.filter(created_at__time__range=[start_time.time(), end_time.time()],
                                  order_status='Ready to ship').count())
                response_data['on_shipping'].append(
                    orders.filter(created_at__time__range=[start_time.time(), end_time.time()],
                                  order_status='On shipping').count())
                response_data['delivered'].append(
                    orders.filter(created_at__time__range=[start_time.time(), end_time.time()],
                                  order_status='Delivered').count())
                response_data['cancelled'].append(
                    orders.filter(created_at__time__range=[start_time.time(), end_time.time()],
                                  order_status='Cancelled').count())
                response_data['returned'].append(
                    orders.filter(created_at__time__range=[start_time.time(), end_time.time()],
                                  order_status='Return').count())

        elif period == 'day':
            for day in range((end_date - start_date).days + 1):
                current_date = start_date + timedelta(days=day)
                response_data['labels'].append(current_date.strftime('%d %b'))
                response_data['accepted'].append(
                    orders.filter(created_at__date=current_date, order_status='Accepted').count())
                response_data['ready_to_ship'].append(
                    orders.filter(created_at__date=current_date, order_status='Ready to ship').count())
                response_data['on_shipping'].append(
                    orders.filter(created_at__date=current_date, order_status='On shipping').count())
                response_data['delivered'].append(
                    orders.filter(created_at__date=current_date, order_status='Delivered').count())
                response_data['cancelled'].append(
                    orders.filter(created_at__date=current_date, order_status='Cancelled').count())
                response_data['returned'].append(
                    orders.filter(created_at__date=current_date, order_status='Return').count())

        elif period == 'month':
            for month in range(start_date.month, end_date.month + 1):
                month_start_date = start_date.replace(day=1, month=month)
                month_end_date = (month_start_date + timedelta(days=32)).replace(day=1) - timedelta(days=1)
                response_data['labels'].append(month_start_date.strftime('%B'))
                response_data['accepted'].append(
                    orders.filter(created_at__date__range=[month_start_date, month_end_date],
                                  order_status='Accepted').count())
                response_data['ready_to_ship'].append(
                    orders.filter(created_at__date__range=[month_start_date, month_end_date],
                                  order_status='Ready to ship').count())
                response_data['on_shipping'].append(
                    orders.filter(created_at__date__range=[month_start_date, month_end_date],
                                  order_status='On shipping').count())
                response_data['delivered'].append(
                    orders.filter(created_at__date__range=[month_start_date, month_end_date],
                                  order_status='Delivered').count())
                response_data['cancelled'].append(
                    orders.filter(created_at__date__range=[month_start_date, month_end_date],
                                  order_status='Cancelled').count())
                response_data['returned'].append(
                    orders.filter(created_at__date__range=[month_start_date, month_end_date],
                                  order_status='Return').count())

        return JsonResponse(response_data)
    return HttpResponse('You are not authorized to view this page', status=403)

@login_required(login_url='login')
def order_status_data(request):
    if request.method == 'POST' and request.user.is_superuser:
        try:
            data = json.loads(request.body)
            start_date = parse_date(data.get('start_date'))
            end_date = parse_date(data.get('end_date'))
        except (ValueError, KeyError):
            return HttpResponse('Invalid data', status=400)

        status_color_mapping = {
            'Accepted': '#80e1c1',
            'Ready to ship': '#f3d676',
            'On shipping': '#f2994a',
            'Delivered': '#4c84ff',
            'Cancelled': '#ff7b7b',
            'Return': '#bb6bd9',
        }

        # Filter orders between start_date and end_date
        queryset = Order.objects.filter(
            created_at__date__gte=start_date,
            created_at__date__lte=end_date
        ).values('order_status').annotate(total=Count('id'))

        labels = []
        data = []
        backgroundColor = []

        for item in queryset:
            status = item['order_status']
            if status in status_color_mapping:
                labels.append(status)
                data.append(item['total'])
                backgroundColor.append(status_color_mapping[status])

        response_data = {
            'labels': labels,
            'data': data,
            'backgroundColor': backgroundColor
        }
        return JsonResponse(response_data)
    return HttpResponse('You are not authorized to view this page', status=403)


@login_required(login_url='login')
def get_payment_by_city(request):
    if request.method == 'POST' and request.user.is_superuser:
        try:
            data = json.loads(request.body)
            start_date = data.get('start_date')
            end_date = data.get('end_date')

            if not start_date or not end_date:
                return JsonResponse({'error': 'Missing start_date or end_date'}, status=400)

            start_date = datetime.strptime(start_date, '%Y-%m-%d')
            end_date = datetime.strptime(end_date, '%Y-%m-%d') + timedelta(days=1)

        except ValueError as e:
            return JsonResponse({'error': str(e)}, status=400)

        city_ids = {
            'R1973756': 'VN-SG',
            'R1875748': 'An Giang',
            'R1873533': 'VN-55',
            'R1904296': 'Bà Rịa–Vũng Tàu',
            'R1902941': 'Bắc Giang',
            'R1903471': 'Bắc Kạn',
            'R1902690': 'Bắc Ninh',
            'R1875968': 'Bến Tre',
            'R1906037': 'Bình Dương',
            'R1889794': 'Bình Định',
            'R1898841': 'Bình Phước',
            'R1904231': 'Bình Thuận',
            'R1873490': 'Cà Mau',
            'R1844412': 'Cao Bằng',
            'R1874283': 'VN-CT',
            'R1891418': 'VN-DN',
            'R1884034': 'Đắk Lắk',
            'R1884042': 'VN-72',
            'R1903340': 'Điện Biên',
            'R1904421': 'Đồng Nai',
            'R1875866': 'Đồng Tháp',
            'R1884018': 'Gia Lai',
            'R1903478': 'Hà Giang',
            'R1902686': 'Hải Dương',
            'R1902682': 'Hải Phòng',
            'R1901010': 'Hà Nam',
            'R1903516': 'VN-HN',
            'R1898458': 'Hà Tĩnh',
            'R1874249': 'Hậu Giang',
            'R1902973': 'Hòa Bình',
            'R1901032': 'Hưng Yên',
            'R1887959': 'Khánh Hòa',
            'R1874471': 'Kiên Giang',
            'R1879515': 'Kon Tum',
            'R1903322': 'Lai Châu',
            'R5522596': 'Lạng Sơn',
            'R1903400': 'Lào Cai',
            'R1885367': 'Lâm Đồng',
            'R1877236': 'Long An',
            'R1901008': 'Nam Định',
            'R1898509': 'Nghệ An',
            'R1900963': 'Ninh Bình',
            'R1886159': 'Ninh Thuận',
            'R1902930': 'Phú Thọ',
            'R1889204': 'Phú Yên',
            'R1896050': 'Quảng Bình',
            'R1891352': 'Quảng Nam',
            'R1890793': 'Quảng Ngãi',
            'R1902947': 'Quảng Ninh',
            'R1895630': 'Quảng Trị',
            'R1873632': 'Sóc Trăng',
            'R1903291': 'Sơn La',
            'R1898961': 'Tây Ninh',
            'R1901019': 'Thái Bình',
            'R1902967': 'Thái Nguyên',
            'R1898590': 'Thanh Hóa',
            'R1891483': 'Thừa Thiên–Huế',
            'R1876011': 'Tiền Giang',
            'R1873642': 'Trà Vinh',
            'R1903418': 'Tuyên Quang',
            'R1875887': 'Vĩnh Long',
            'R1902889': 'Vĩnh Phúc',
            'R1903199': 'Yên Bái'
        }

        totals = {name: 0 for name in city_ids.values()}

        payments = Payment.objects.filter(
            order__created_at__range=[start_date, end_date],
            order__city__in=city_ids.keys(),
            order__is_ordered=True
        ).values('order__city', 'amount_paid')

        for payment in payments:
            city_name = city_ids.get(payment['order__city'])
            if city_name:
                totals[city_name] += payment['amount_paid']

        result = [['City', 'Revenue']]
        for city, total in totals.items():
            result.append([city, total])

        return JsonResponse(result, safe=False)
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def top_purchase(request):
    if request.method == 'POST' and  request.user.is_superuser:
        try:
            data = json.loads(request.body)
            start_date = data.get('start_date')
            end_date = data.get('end_date')

            if not start_date or not end_date:
                return JsonResponse({'error': 'Missing start_date or end_date'}, status=400)

            start_date = datetime.strptime(start_date, '%Y-%m-%d')
            end_date = datetime.strptime(end_date, '%Y-%m-%d') + timedelta(days=1)

        except ValueError as e:
            return JsonResponse({'error': str(e)}, status=400)

        payments = Payment.objects.filter(
            order__created_at__range=[start_date, end_date],
            order__is_ordered=True
        ).values('order__city').annotate(
            total_paid=Sum('amount_paid')
        ).order_by('-total_paid')[:3]

        response = {
            'labels': [get_locations(payment['order__city'], '') for payment in payments],
            'data': [round(payment['total_paid'], 2) for payment in payments]
        }

        return JsonResponse(response)

    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def recent_orders(request):
    if request.method == 'POST' and request.user.is_superuser:
        orders = Order.objects.prefetch_related('orderproduct_set__product').order_by('-created_at')[:5]
        orders_data = []
        current_site = get_current_site(request)
        for order in orders:
            products_info = [
                {
                    'product_name': op.product.product_name,
                    'quantity': op.quantity,
                    'product_price': op.product_price,
                    "product_url": f"http://{current_site}/store/{op.product.category_main.slug}/{op.product.sub_category.slug}/{op.product.slug}"
                }
                for op in order.orderproduct_set.all()
            ]
            order_dict = {
                "order_id": order.id,
                "order_number": order.order_number,
                "order_date": order.created_at.strftime("%b %d, %Y"),
                "order_cost": order.order_total,
                "status": order.order_status,
                "products": products_info,
                'url_view': f"http://{current_site}/accounts/order_details/{order.order_number}",
                'url_edit': f"http://{current_site}/admin/order_update/{str(order.id)}",
            }
            orders_data.append(order_dict)
        return JsonResponse(orders_data, safe=False)
    return HttpResponse('You are not authorized to view this page')

@login_required(login_url='login')
def order_list(request):
    if request.user.is_superuser:
        order_list_all = Order.objects.select_related('user', 'payment').prefetch_related('orderproduct_set').order_by('-created_at')
        paginator = Paginator(order_list_all, 10) # Show 10 orders per page

        page_number = request.GET.get('page')
        try:
            page_obj = paginator.get_page(page_number)
        except PageNotAnInteger:
            page_obj = paginator.page(1)
        except EmptyPage:
            page_obj = paginator.page(paginator.num_pages)

        # order_products might not be needed if accessed via page_obj.object_list in template
        # order_products = OrderProduct.objects.filter(order__in=page_obj.object_list)

        context = {
            'page_obj': page_obj, # Pass the page object for orders
            # 'order_products': order_products, # Pass only if needed separately
        }
        return render(request, 'adminApp/Orders/order_list.html', context)
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def order_update(request, pk):
    if request.user.is_superuser:
        order = Order.objects.get(pk=pk)
        form = OrderUpdateForm(instance=order)
        if request.method == 'POST':
            form = OrderUpdateForm(request.POST, instance=order)
            if form.is_valid():
                form.save()
                return redirect('order_list')
            else:
                print(form.errors)
        context = {
            'form': form,
            'order': order,
        }
        return render(request, 'adminApp/orders/order_update.html', context)
    return HttpResponse('You are not authorized to view this page')


@login_required(login_url='login')
def user_profile_api(request, pk):

    if request.user.is_superuser:
        try:
            user = Account.objects.get(pk=pk)
            profile = UserProfile.objects.get(user=user)
            
            user_data = {
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'phone_number': user.phone_number,
                'full_name': user.full_name,
                'date_joined': user.date_joined,
                'last_login': user.last_login,
                'is_active': user.is_active,
                'is_admin': user.is_admin,
                'is_login': user.is_login,
            }
            
            profile_data = {
                'id': profile.id,
                'date_of_birth': profile.date_of_birth,
                'sex': profile.sex,
                'road': profile.road,
                'ward': profile.ward,
                'district': profile.district,
                'city': profile.city,
                'ward_name': profile.ward_name(),
                'district_name': profile.district_name(),
                'city_name': profile.city_name(),
                'full_address': profile.full_address(),
                'profile_picture': profile.profile_picture,
                'bio': profile.bio,
            }
            
            orders = Order.objects.filter(user=user).order_by('-created_at')
            orders_data = []
            for order in orders:
                order_items = OrderProduct.objects.filter(order=order)
                items_data = []
                for item in order_items:
                    items_data.append({
                        'product_name': item.product.product_name,
                        'quantity': item.quantity,
                        'price': float(item.product_price),
                        'subtotal': float(item.product_price * item.quantity),
                    })
                
                orders_data.append({
                    'id': order.id,
                    'order_number': order.order_number,
                    'status': order.status,
                    'created_at': order.created_at,
                    'updated_at': order.updated_at,
                    'total': float(order.order_total),
                    'tax': float(order.tax),
                    'items': items_data,
                    'payment_method': order.payment.payment_method if hasattr(order, 'payment') else None,
                    'payment_status': order.payment.status if hasattr(order, 'payment') else None,
                })
            
            wishlist_items = Wishlist.objects.filter(user=user).order_by('-created_at')
            wishlist_data = []
            for item in wishlist_items:
                wishlist_data.append({
                    'id': item.id,
                    'product_id': item.product.id,
                    'product_name': item.product.product_name,
                    'price': float(item.product.price),
                    'image': item.product.images,
                    'created_at': item.created_at,
                })
            
            events = EventUser.objects.filter(user=user).order_by('-event_timestamp')
            events_data = []
            for event in events:
                event_data = {
                    'id': event.id,
                    'event_type': event.event_type,
                    'event_timestamp': event.event_timestamp,
                    'frequency': event.frequency,
                }
                if event.product:
                    event_data['product'] = {
                        'id': event.product.id,
                        'name': event.product.product_name,
                    }
                events_data.append(event_data)
            
            # Thống kê
            view_count = EventUser.objects.filter(user=user, event_type='view').count()
            cart_count = EventUser.objects.filter(user=user, event_type='cart').count()
            pay_count = EventUser.objects.filter(user=user, event_type='pay').count()
            login_count = EventUser.objects.filter(user=user, event_type='login').count()
            
            total_spent = sum(float(order.order_total) for order in orders)
            
            stats = {
                'view_count': view_count,
                'cart_count': cart_count,
                'pay_count': pay_count,
                'login_count': login_count,
                'total_spent': total_spent,
                'order_count': orders.count(),
                'wishlist_count': wishlist_items.count(),
                'last_view': events.filter(event_type='view').first().event_timestamp if events.filter(event_type='view').exists() else None,
                'last_login': events.filter(event_type='login').first().event_timestamp if events.filter(event_type='login').exists() else None,
            }
            
            response_data = {
                'user': user_data,
                'profile': profile_data,
                'orders': orders_data,
                'wishlist': wishlist_data,
                'events': events_data,
                'stats': stats,
            }
            
            return JsonResponse(response_data)
            
        except Account.DoesNotExist:
            return JsonResponse({'error': 'User not found'}, status=404)
        except UserProfile.DoesNotExist:
            return JsonResponse({'error': 'User profile not found'}, status=404)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'You are not authorized to access this information'}, status=403)


AES_KEY = base64.urlsafe_b64decode(config('AES_KEY'))


def encrypt(data):
    iv = os.urandom(16)
    cipher = Cipher(algorithms.AES(AES_KEY), modes.CBC(iv), backend=default_backend())
    encryptor = cipher.encryptor()

    padder = padding.PKCS7(128).padder()
    padded_data = padder.update(data.encode()) + padder.finalize()

    encrypted_data = encryptor.update(padded_data) + encryptor.finalize()
    return base64.urlsafe_b64encode(iv + encrypted_data).decode('utf-8')


def decrypt(encrypted_data):
    encrypted_data = base64.urlsafe_b64decode(encrypted_data)
    iv = encrypted_data[:16]
    encrypted_message = encrypted_data[16:]

    cipher = Cipher(algorithms.AES(AES_KEY), modes.CBC(iv), backend=default_backend())
    decryptor = cipher.decryptor()

    decrypted_padded_message = decryptor.update(encrypted_message) + decryptor.finalize()

    unpadder = padding.PKCS7(128).unpadder()
    decrypted_message = unpadder.update(decrypted_padded_message) + unpadder.finalize()

    return decrypted_message.decode('utf-8')


@login_required(login_url='login')
def get_folders(request):
    if request.method == 'GET' and request.user.is_superuser:
        try:
            response = requests.get('http://127.0.0.1:5000/get-folders')

            if response.status_code == 200:
                data = response.json()
                folders = []

                for encrypted_folder in data.get('folders', []):
                    folder_name = decrypt(encrypted_folder)

                    is_imported = FolderEvent.objects.filter(name=folder_name).exists()

                    folders.append({
                        'name': folder_name,
                        'imported': is_imported
                    })

                return JsonResponse({"status": "success", "folders": folders})
            else:
                return JsonResponse({
                    "status": "error",
                    "message": f"ETL service returned status code {response.status_code}"
                }, status=500)

        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)}, status=500)

    return JsonResponse({"status": "error", "message": "Unauthorized or invalid request"}, status=403)


@login_required(login_url='login')
def process_folder(request):
    if request.method == 'POST' and request.user.is_superuser:
        try:
            data = json.loads(request.body)
            folder_name = data.get('folder_name')

            if not folder_name:
                return JsonResponse({"status": "error", "message": "Folder name is required"}, status=400)

            encrypted_data = encrypt(json.dumps({"folder_name": folder_name}))

            response = requests.post(
                'http://127.0.0.1:5000/process-folder',
                json={"data": encrypted_data}
            )

            if response.status_code == 200:
                result = response.json()

                if result.get('status') == 'success':
                    FolderEvent.objects.create(name=folder_name)

                    return JsonResponse({
                        "status": "success",
                        "message": "Folder imported successfully"
                    })
                else:
                    return JsonResponse({
                        "status": "error",
                        "message": result.get('message', 'Failed to process folder')
                    })
            else:
                return JsonResponse({
                    "status": "error",
                    "message": f"ETL service returned status code {response.status_code}"
                }, status=500)

        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)}, status=500)

    return JsonResponse({"status": "error", "message": "Unauthorized or invalid request"}, status=403)
